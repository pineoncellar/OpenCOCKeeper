# -*- coding: utf-8 -*-
"""
@File     :   conftest.py
@Desc     :   pytest 共享夹具：独立临时库、测试段 world_id、FakeLLM、全局单例隔离
@Note     :   数据隔离双保险——每个测试用 tmp_path 独立 SQLite 库（物理隔离），
             world_id 固定落在 900 测试段（逻辑隔离），并在测试期间清空全局
             db/storage 单例指针，杜绝误触真实 data/app.db
"""

from __future__ import annotations

import pytest

from src import llm as llm_module
from src.core import config as cfg_module
from src.core import db as db_module
from src.core import ids
from src.core.db import Database
from src.llm import client as llm_client
from src.llm.fake import FakeLLM
from src.storage import storage as storage_module
from src.storage.storage import Storage

# 每个测试由 autouse 夹具重建的全新 FakeLLM，测试内经 fake_llm fixture 配置
_fake_llm: FakeLLM | None = None

# _tmp_modules 夹具在每个测试的临时目录创建的可绑定模组文件名
TEST_MODULE_NAME = "test_module.docx"


def pytest_configure(config):
    config.addinivalue_line(
        "markers", "real_llm: 该测试使用真实 LLM 客户端（跳过 FakeLLM 注入）"
    )


@pytest.fixture
def tmp_db(tmp_path):
    """每个测试独立的 SQLite 库，物理隔离真实 data/app.db。"""
    return Database(tmp_path / "test_app.db")


@pytest.fixture
def storage(tmp_db):
    """基于独立库的存储门面（自动迁移到最新 schema）。"""
    return Storage(db=tmp_db)


@pytest.fixture(autouse=True)
def _tmp_modules(tmp_path, monkeypatch):
    """每个测试独立的临时模组目录（含 test_module.docx 与对应划界缓存），并替换模块层 MODULES_DIR。

    划界缓存预写为单章"测试模组"，让 build_index 走缓存路径不触发 LLM；
    世界绑定强校验要求模组文件必须存在，测试世界统一绑定该文件。
    """
    import json
    from src.module import loader as module_loader
    from docx import Document
    d = tmp_path / "modules"
    d.mkdir()
    p = d / TEST_MODULE_NAME
    doc = Document()
    doc.add_heading("测试模组", level=1)
    doc.add_paragraph("这是测试模组的正文内容。")
    doc.save(p)
    # 预写划界缓存（单章），使 build_index 走缓存不触发 LLM 划界  # 状态：缓存预置
    cache = d / ".cache"
    cache.mkdir()
    (cache / f"{TEST_MODULE_NAME}.json").write_text(
        json.dumps(
            {
                "source_mtime": p.stat().st_mtime,
                "sections": [{"title": "测试模组", "start_anchor": "测试模组"}],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(module_loader, "MODULES_DIR", d)
    return d


@pytest.fixture(autouse=True)
def _tmp_trace_store(tmp_path, monkeypatch):
    """每个测试独立的 trace 目录并重置 TraceStore 单例，杜绝写真实 logs/traces。

    TraceBus.publish 内部会调 get_trace_store().append 落盘，若不做隔离，
    任何走 run_tool_loop / narrate 的测试都会污染工作区 logs/traces。
    """
    from src.webui import trace_store as trace_store_module

    d = tmp_path / "traces"
    d.mkdir()
    monkeypatch.setattr(trace_store_module, "TRACE_DIR", d)
    monkeypatch.setattr(trace_store_module, "_trace_store", None)
    return d


@pytest.fixture
def world_id(storage, _tmp_modules):
    """测试段 world_id（900 段），逻辑上远离真实世界的 world_001.. 段。"""
    wid = ids.make_world_id(900, "test")
    storage.ensure_world(wid, module_name=TEST_MODULE_NAME)
    return wid


@pytest.fixture(autouse=True)
def _auto_fake_llm(monkeypatch, request):
    """默认注入 FakeLLM：测试中调用 call_llm / call_llm_stream / ask_llm 一律走 fake，
    不触真实网络与密钥；标 @pytest.mark.real_llm 的测试跳过注入，走真实实现。
    """
    global _fake_llm
    if "real_llm" in request.keywords:
        yield
        return
    _fake_llm = FakeLLM()  # 状态：每测试全新实例，responses/calls 天然隔离
    # 同时替换 client 模块与包命名空间两处入口，覆盖不同导入方式
    monkeypatch.setattr(llm_client, "call_llm", _fake_llm.call)
    monkeypatch.setattr(llm_client, "call_llm_stream", _fake_llm.call_stream)
    monkeypatch.setattr(llm_client, "ask_llm", _fake_llm.ask)
    monkeypatch.setattr(llm_module, "call_llm", _fake_llm.call)
    monkeypatch.setattr(llm_module, "call_llm_stream", _fake_llm.call_stream)
    monkeypatch.setattr(llm_module, "ask_llm", _fake_llm.ask)
    yield


@pytest.fixture
def fake_llm(_auto_fake_llm):
    """当前测试注入的那个 FakeLLM 实例，测试内可设预设响应 / 读调用历史。"""
    assert _fake_llm is not None
    return _fake_llm


@pytest.fixture(autouse=True)
def _isolate_singletons():
    """备份并清空全局 db/storage 单例，防止测试误触真实库；结束恢复。"""
    old_db, old_storage = db_module._db, storage_module._storage
    db_module._db, storage_module._storage = None, None
    yield
    db_module._db, storage_module._storage = old_db, old_storage
