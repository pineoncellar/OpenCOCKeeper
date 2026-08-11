# -*- coding: utf-8 -*-
"""
@File     :   test_retrieval.py
@Desc     :   模组检索层测试：锚点切片（原文逐字对齐）、LLM 划界（FakeLLM）、索引缓存、检索门面
@Note     :   依赖 conftest 的 _tmp_modules autouse 夹具（临时模组目录 + 划界缓存预置）；
             LLM 划界用 FakeLLM 预设响应，与真实网络解耦；PDF 用例无 pypdf 时跳过
"""

from __future__ import annotations

import json

import pytest

from src.core.ids import make_world_id
from src.retrieval import build_index, clear_cache, search_module
from src.retrieval import structure
from src.retrieval.sections import split_sections

# conftest 预置的单章测试模组（标题"测试模组"）
TEST_MODULE = "test_module.docx"

# 测试用丰富纯文本：多章节 + 专有名词，覆盖锚点切分与检索
RICH_MODULE = "rich.docx"
RICH_TEXT = (
    "前言部分，调查员尚未抵达。\n\n"
    "地点1：瑞德农庄\n\n"
    "调查员在通往农庄的小路上遇到了爱丽丝·卡罗尔。\n\n"
    "厨房\n\n"
    "厨房里弥漫着腐肉的气味，灶台上放着一本沾血的日记。\n\n"
    "地点3：杰克房间\n\n"
    "房间的墙壁上布满抓痕，杰克的书桌上摆着失踪当天的报纸。"
)
RICH_METAS = [
    {"title": "地点1：瑞德农庄", "start_anchor": "地点1：瑞德农庄"},
    {"title": "厨房", "start_anchor": "厨房"},
    {"title": "地点3：杰克房间", "start_anchor": "地点3：杰克房间"},
]


@pytest.fixture
def rich_module(_tmp_modules):
    """构造多章节 Word 模组 + 对应划界缓存，供检索测试使用。"""
    from docx import Document
    d = _tmp_modules
    p = d / RICH_MODULE
    doc = Document()
    doc.add_heading("地点1：瑞德农庄", level=1)
    doc.add_paragraph("调查员在通往农庄的小路上遇到了爱丽丝·卡罗尔。")
    doc.add_heading("厨房", level=1)
    doc.add_paragraph("厨房里弥漫着腐肉的气味，灶台上放着一本沾血的日记。")
    doc.add_heading("地点3：杰克房间", level=1)
    doc.add_paragraph("房间的墙壁上布满抓痕，杰克的书桌上摆着失踪当天的报纸。")
    doc.save(p)
    cache = d / ".cache"
    cache.mkdir(exist_ok=True)
    (cache / f"{RICH_MODULE}.json").write_text(
        json.dumps({"source_mtime": p.stat().st_mtime, "sections": RICH_METAS}, ensure_ascii=False),
        encoding="utf-8",
    )
    clear_cache()
    return RICH_MODULE


# ====================================================================
# 锚点切片（原文逐字对齐，硬性保证）
# ====================================================================


def test_split_sections_anchors_verbatim():
    sections = split_sections(RICH_TEXT, RICH_METAS)
    titles = [s.title for s in sections]
    assert titles == ["引言", "地点1：瑞德农庄", "厨房", "地点3：杰克房间"]
    # 引言 = 首个锚点之前的原文（含分隔换行，逐字保留）
    assert sections[0].content == "前言部分，调查员尚未抵达。\n\n"
    # 各节 content 为原文切片：含标题行与正文（物理切片从锚点起到下一锚点前）
    kitchen = sections[2]
    assert "厨房" in kitchen.content
    assert "厨房里弥漫着腐肉的气味，灶台上放着一本沾血的日记。" in kitchen.content
    # 拼接回原文，验证 100% 物理对齐
    assert "".join(s.content for s in sections) == RICH_TEXT


def test_split_sections_no_anchor_falls_back():
    sections = split_sections(RICH_TEXT, [])
    assert len(sections) == 1
    assert sections[0].content == RICH_TEXT
    assert sections[0].title == "全文"


def test_split_sections_missing_anchor_skipped():
    metas = RICH_METAS + [{"title": "幽灵章节", "start_anchor": "从未出现的锚点"}]
    sections = split_sections(RICH_TEXT, metas)
    assert "幽灵章节" not in [s.title for s in sections]


def test_split_sections_duplicate_anchor_dedup():
    metas = RICH_METAS + [{"title": "重复", "start_anchor": "厨房"}]
    sections = split_sections(RICH_TEXT, metas)
    assert [s.title for s in sections].count("厨房") == 1


# ====================================================================
# LLM 划界（FakeLLM，与真实网络解耦）
# ====================================================================


def test_structure_parses_json(fake_llm):
    fake_llm.set_response(
        "fast", json.dumps({"sections": [{"title": "厨房", "start_anchor": "厨房"}]}, ensure_ascii=False),
    )
    assert structure.delineate_sync("前言\n\n厨房\n\n正文。") == [
        {"title": "厨房", "start_anchor": "厨房"},
    ]


def test_structure_non_json_falls_back(fake_llm):
    fake_llm.set_response("fast", "抱歉，我无法处理。")
    assert structure.delineate_sync("随便什么文本") == []


def test_structure_empty_sections(fake_llm):
    fake_llm.set_response("fast", json.dumps({"sections": []}))
    assert structure.delineate_sync("随便什么文本") == []


def test_structure_chunk_and_merge(fake_llm, monkeypatch):
    # 极小分块强制多块调用，重叠锚点合并去重
    monkeypatch.setattr(
        structure, "_settings",
        lambda: {"enabled": True, "tier": "fast", "chunk_chars": 30},
    )
    fake_llm.set_response(
        "fast", json.dumps({"sections": [{"title": "地点1", "start_anchor": "地点1"}]}, ensure_ascii=False),
    )
    metas = structure.delineate_sync("地点1\n\n内容……" * 8)
    assert metas == [{"title": "地点1", "start_anchor": "地点1"}]
    assert len(fake_llm.calls) >= 2  # 触发分块多次调用


# ====================================================================
# 索引与缓存（JSON 划界缓存 + mtime 失效）
# ====================================================================


def test_index_cache_hit_no_llm(fake_llm):
    # conftest 预置划界缓存，build_index 不应触发 LLM
    idx = build_index(TEST_MODULE)
    assert fake_llm.calls == []
    assert idx.sections[0].title == "测试模组"
    assert "测试模组" in idx.sections[0].content


def test_index_rebuilds_on_mtime_change(_tmp_modules, fake_llm):
    import os
    from docx import Document
    p = _tmp_modules / TEST_MODULE
    doc = Document()
    doc.add_heading("新标题", level=1)
    doc.add_paragraph("新正文。")
    doc.save(p)
    os.utime(p, (p.stat().st_atime, p.stat().st_mtime + 1))
    fake_llm.set_response(
        "fast", json.dumps({"sections": [{"title": "新标题", "start_anchor": "新标题"}]}, ensure_ascii=False),
    )
    clear_cache()
    idx = build_index(TEST_MODULE)
    assert idx.sections[0].title == "新标题"
    assert fake_llm.calls  # 缓存失效后触发划界


# ====================================================================
# 检索门面：标题加权 / BM25 / top_k / 世界绑定
# ====================================================================


def test_search_title_match_ranked_first(rich_module):
    hits = search_module("厨房", module_name=rich_module, top_k=3)
    assert hits and hits[0].title_match
    assert hits[0].section.title == "厨房"


def test_search_bm25_fulltext_hit(rich_module):
    # "爱丽丝"只出现在正文，标题不含 → 全文命中
    hits = search_module("爱丽丝 卡罗尔", module_name=rich_module, top_k=3)
    assert hits[0].section.title == "地点1：瑞德农庄"
    assert not hits[0].title_match


def test_search_default_top_k_is_two(rich_module):
    hits = search_module("房间 报纸 日记", module_name=rich_module)
    assert len(hits) <= 2


def test_search_no_target_returns_empty():
    assert search_module("厨房") == []


def test_search_by_world_binding(storage, world_id, monkeypatch, tmp_db):
    # conftest 的 world_id 已绑定 test_module.docx（标题"测试模组"）；
    # 把 search_module 内部读世界绑定的 get_db 指向测试库
    from src.retrieval import search as search_mod
    monkeypatch.setattr(search_mod, "get_db", lambda: tmp_db)
    hits = search_module("测试模组", world_id=world_id, top_k=3)
    assert hits and hits[0].section.title == "测试模组"


def test_search_unknown_world_returns_empty(storage, monkeypatch, tmp_db):
    from src.retrieval import search as search_mod
    monkeypatch.setattr(search_mod, "get_db", lambda: tmp_db)
    assert search_module("厨房", world_id=make_world_id(900, "nope")) == []


# ====================================================================
# PDF 纯文本提取（无 pypdf 时跳过）
# ====================================================================

pypdf = pytest.importorskip("pypdf")


def test_pdf_read_pure_text(_tmp_modules):
    from pypdf import PdfWriter
    from src.module.reader import read_module
    path = _tmp_modules / "demo.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    with open(path, "wb") as f:
        writer.write(f)
    # 空白 PDF 无文字层，read_module 返回空串且不崩溃
    assert read_module("demo.pdf") == ""


# ====================================================================
# 异步检索（运行中事件循环内无缓存路径，回归 delineate asyncio.run 冲突）
# ====================================================================


async def test_build_index_async_no_cache_within_event_loop(tmp_path, monkeypatch, fake_llm):
    """无划界缓存的模组在运行中事件循环内异步建索引：走 await delineate，
    不再触发 delineate_sync 的 asyncio.run 冲突（Opening Agent 首轮检索 bug 回归）。"""
    from docx import Document
    from src.module import loader as module_loader
    from src.retrieval import build_index_async, clear_cache, search_module_async

    d = tmp_path / "mods"
    d.mkdir()
    p = d / "no_cache.docx"
    doc = Document()
    doc.add_heading("无缓存模组", level=1)
    doc.add_paragraph("雨后的清晨，委托人叩响了事务所的门。")
    doc.save(p)
    monkeypatch.setattr(module_loader, "MODULES_DIR", d)
    clear_cache("no_cache.docx")  # 状态：清内存索引，确保走无缓存划界路径

    # fake_llm 划界返回 "fake-ok"（不可解析 JSON）→ 空 metas → 整篇一段兜底，不崩
    idx = await build_index_async("no_cache.docx")
    assert idx.sections and idx.sections[0].title == "全文"

    # 检索门面异步版同样可用（运行中循环内 await 划界/检索）
    hits = await search_module_async("委托人", module_name="no_cache.docx")
    assert hits and "委托人" in hits[0].section.content

